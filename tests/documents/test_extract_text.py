from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from PIL import Image
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, EncodedStreamObject, NameObject

from src.documents.extract_text import UnsupportedDocument, extract_document


def test_extracts_utf8_text() -> None:
    result = extract_document("resume.txt", "姓名：林晓\n院校：示例大学".encode())
    assert "林晓" in result.text
    assert result.mime_type == "text/plain"


def test_rejects_executable_content_disguised_as_text() -> None:
    with pytest.raises(UnsupportedDocument):
        extract_document("resume.txt", b"MZ")


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("姓名：陈澜")
    document.add_paragraph("   ")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "技能"
    table.cell(0, 1).text = "Python"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _text_pdf_bytes() -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=300)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    resources = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    page[NameObject("/Resources")] = resources
    stream = DecodedStreamObject()
    stream.set_data(b"BT /F1 12 Tf 20 200 Td (Anonymous Resume) Tj ET")
    page[NameObject("/Contents")] = writer._add_object(stream)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def _form_pdf_bytes() -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=300)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    form = DecodedStreamObject()
    form.set_data(b"BT /F1 12 Tf 20 200 Td (Nested Resume) Tj ET " + b"%" + b"A" * 100)
    form[NameObject("/Type")] = NameObject("/XObject")
    form[NameObject("/Subtype")] = NameObject("/Form")
    form[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})}
    )
    form_ref = writer._add_object(form)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/XObject"): DictionaryObject({NameObject("/Fm1"): form_ref})}
    )
    direct = DecodedStreamObject()
    direct.set_data(b"q /Fm1 Do Q")
    page[NameObject("/Contents")] = writer._add_object(direct)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def _cyclic_form_pdf_bytes() -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=10, height=10)
    first = DecodedStreamObject()
    first.set_data(b"q /Fm2 Do Q")
    first[NameObject("/Subtype")] = NameObject("/Form")
    second = DecodedStreamObject()
    second.set_data(b"q /Fm1 Do Q")
    second[NameObject("/Subtype")] = NameObject("/Form")
    first_ref = writer._add_object(first)
    second_ref = writer._add_object(second)
    first[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/XObject"): DictionaryObject({NameObject("/Fm2"): second_ref})}
    )
    second[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/XObject"): DictionaryObject({NameObject("/Fm1"): first_ref})}
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/XObject"): DictionaryObject({NameObject("/Fm1"): first_ref})}
    )
    direct = DecodedStreamObject()
    direct.set_data(b"q /Fm1 Do Q")
    page[NameObject("/Contents")] = writer._add_object(direct)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def _image_bytes(image_format: str, size: tuple[int, int] = (4, 3)) -> bytes:
    output = BytesIO()
    Image.new("RGB", size, "white").save(output, format=image_format)
    return output.getvalue()


def _append_zip_member(source: bytes, name: str, data: bytes) -> bytes:
    output = BytesIO(source)
    with ZipFile(output, "a", compression=ZIP_DEFLATED) as archive:
        archive.writestr(name, data)
    return output.getvalue()


def _replace_zip_member(source: bytes, name: str, data: bytes) -> bytes:
    output = BytesIO()
    with (
        ZipFile(BytesIO(source)) as original,
        ZipFile(output, "w", compression=ZIP_DEFLATED) as rewritten,
    ):
        for member in original.infolist():
            rewritten.writestr(
                member.filename, data if member.filename == name else original.read(member)
            )
    return output.getvalue()


def test_extracts_docx_paragraphs_and_tables() -> None:
    result = extract_document("resume.docx", _docx_bytes())
    assert (
        result.mime_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert result.text.splitlines() == ["姓名：陈澜", "技能", "Python"]
    assert result.image_bytes == ()


def test_extracts_text_pdf() -> None:
    result = extract_document("resume.pdf", _text_pdf_bytes())
    assert "Anonymous Resume" in result.text
    assert result.mime_type == "application/pdf"


@pytest.mark.parametrize(
    ("filename", "image_format", "mime_type"),
    [("photo.PNG", "PNG", "image/png"), ("photo.JpEg", "JPEG", "image/jpeg")],
)
def test_validates_images(filename: str, image_format: str, mime_type: str) -> None:
    result = extract_document(filename, _image_bytes(image_format))
    assert result.text == ""
    assert result.mime_type == mime_type
    assert len(result.image_bytes) == 1
    Image.open(BytesIO(result.image_bytes[0])).verify()


def test_supports_uppercase_text_extension() -> None:
    assert extract_document("RESUME.TXT", b"candidate").text == "candidate"


@pytest.mark.parametrize(
    ("filename", "content"),
    [("empty.txt", b""), ("broken.docx", b"PK\x03\x04broken"), ("broken.pdf", b"%PDF-broken")],
)
def test_rejects_empty_or_corrupt_documents(filename: str, content: bytes) -> None:
    with pytest.raises(UnsupportedDocument, match="空|损坏|无法"):
        extract_document(filename, content)


def test_rejects_oversized_file(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("src.documents.extract_text.MAX_FILE_BYTES", 3)
    with pytest.raises(UnsupportedDocument, match="大小"):
        extract_document("resume.txt", b"four")


def test_rejects_invalid_utf8() -> None:
    with pytest.raises(UnsupportedDocument, match="UTF-8"):
        extract_document("resume.txt", b"\xff")


def test_rejects_excessive_extracted_text(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("src.documents.extract_text.MAX_TEXT_CHARS", 3)
    with pytest.raises(UnsupportedDocument, match="文本"):
        extract_document("resume.txt", b"four")


def test_rejects_encrypted_pdf() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=10, height=10)
    writer.encrypt("secret")
    output = BytesIO()
    writer.write(output)
    with pytest.raises(UnsupportedDocument, match="加密|密码"):
        extract_document("resume.pdf", output.getvalue())


def test_image_only_pdf_reports_vision_conversion_unavailable() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=10, height=10)
    output = BytesIO()
    writer.write(output)
    with pytest.raises(UnsupportedDocument, match="OCR|视觉|图像"):
        extract_document("scan.pdf", output.getvalue())


def test_rejects_binary_extension_mismatch() -> None:
    with pytest.raises(UnsupportedDocument, match="格式|扩展名"):
        extract_document("resume.pdf", _image_bytes("PNG"))


def test_rejects_supported_binary_disguised_as_text() -> None:
    with pytest.raises(UnsupportedDocument, match="格式|扩展名"):
        extract_document("resume.txt", _text_pdf_bytes())


def test_rejects_image_pixel_limit(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("src.documents.extract_text.MAX_IMAGE_PIXELS", 5)
    with pytest.raises(UnsupportedDocument, match="像素|尺寸"):
        extract_document("resume.png", _image_bytes("PNG", (3, 2)))


def test_rejects_control_heavy_utf8_binary() -> None:
    with pytest.raises(UnsupportedDocument, match="二进制|控制字符"):
        extract_document("resume.txt", b"name\x00\x01\x02\x03\x04")


@pytest.mark.parametrize(
    ("constant", "value", "message"),
    [
        ("MAX_DOCX_MEMBERS", 1, "条目|成员"),
        ("MAX_DOCX_MEMBER_BYTES", 10, "成员|解压"),
        ("MAX_DOCX_UNCOMPRESSED_BYTES", 20, "解压|总"),
    ],
)
def test_rejects_docx_zip_budgets(
    monkeypatch: pytest.MonkeyPatch, constant: str, value: int, message: str
) -> None:
    monkeypatch.setattr(f"src.documents.extract_text.{constant}", value)
    with pytest.raises(UnsupportedDocument, match=message):
        extract_document("resume.docx", _docx_bytes())


def test_rejects_docx_suspicious_compression_ratio(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("src.documents.extract_text.MAX_DOCX_COMPRESSION_RATIO", 2)
    content = _append_zip_member(_docx_bytes(), "word/padding.bin", b"A" * 2_000)
    with pytest.raises(UnsupportedDocument, match="压缩比"):
        extract_document("resume.docx", content)


def test_rejects_docx_duplicate_member() -> None:
    with pytest.warns(UserWarning, match="Duplicate name"):
        content = _append_zip_member(_docx_bytes(), "word/document.xml", b"duplicate")
    with pytest.raises(UnsupportedDocument, match="重复"):
        extract_document("resume.docx", content)


def test_rejects_docx_unsafe_member_name() -> None:
    content = _append_zip_member(_docx_bytes(), "../escape.xml", b"unsafe")
    with pytest.raises(UnsupportedDocument, match="路径|名称"):
        extract_document("resume.docx", content)


def test_rejects_docx_encrypted_member_flag() -> None:
    content = bytearray(_docx_bytes())
    central_header = content.find(b"PK\x01\x02")
    assert central_header >= 0
    flags_offset = central_header + 8
    flags = int.from_bytes(content[flags_offset : flags_offset + 2], "little") | 0x1
    content[flags_offset : flags_offset + 2] = flags.to_bytes(2, "little")
    with pytest.raises(UnsupportedDocument, match="加密"):
        extract_document("resume.docx", bytes(content))


def test_translates_malformed_docx_xml() -> None:
    content = _replace_zip_member(_docx_bytes(), "word/document.xml", b"<broken")
    with pytest.raises(UnsupportedDocument, match="DOCX.*损坏|DOCX.*无法"):
        extract_document("resume.docx", content)


def test_rejects_excessive_docx_text_incrementally(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("src.documents.extract_text.MAX_TEXT_CHARS", 3)
    with pytest.raises(UnsupportedDocument, match="文本"):
        extract_document("resume.docx", _docx_bytes())


def test_rejects_excessive_pdf_pages(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("src.documents.extract_text.MAX_PDF_PAGES", 1)
    writer = PdfWriter()
    writer.add_blank_page(width=10, height=10)
    writer.add_blank_page(width=10, height=10)
    output = BytesIO()
    writer.write(output)
    with pytest.raises(UnsupportedDocument, match="页数"):
        extract_document("resume.pdf", output.getvalue())


def test_rejects_pdf_decoded_stream_budget(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("src.documents.extract_text.MAX_PDF_CONTENT_STREAM_BYTES", 10)
    with pytest.raises(UnsupportedDocument, match="内容流|解码"):
        extract_document("resume.pdf", _text_pdf_bytes())


def test_counts_nested_form_xobject_toward_pdf_stream_budget(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr("src.documents.extract_text.MAX_PDF_CONTENT_STREAM_BYTES", 50)
    with pytest.raises(UnsupportedDocument, match="内容流|解码"):
        extract_document("resume.pdf", _form_pdf_bytes())


def test_caps_pdf_resource_objects(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("src.documents.extract_text.MAX_PDF_RESOURCE_OBJECTS", 0)
    with pytest.raises(UnsupportedDocument, match="资源|对象"):
        extract_document("resume.pdf", _form_pdf_bytes())


def test_rejects_cyclic_pdf_form_resources() -> None:
    with pytest.raises(UnsupportedDocument, match="循环|资源"):
        extract_document("resume.pdf", _cyclic_form_pdf_bytes())


def test_rejects_excessive_pdf_text_incrementally(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("src.documents.extract_text.MAX_TEXT_CHARS", 5)
    with pytest.raises(UnsupportedDocument, match="文本"):
        extract_document("resume.pdf", _text_pdf_bytes())


def test_translates_unsupported_pdf_stream_filter() -> None:
    writer = PdfWriter()
    page = writer.add_blank_page(width=10, height=10)
    stream = EncodedStreamObject()
    stream._data = b"invalid"
    stream[NameObject("/Filter")] = NameObject("/UnsupportedFilter")
    page[NameObject("/Contents")] = writer._add_object(stream)
    output = BytesIO()
    writer.write(output)
    with pytest.raises(UnsupportedDocument, match="PDF.*损坏|PDF.*解码"):
        extract_document("resume.pdf", output.getvalue())


def test_rejects_reencoded_image_over_output_limit(monkeypatch: pytest.MonkeyPatch) -> None:
    output = BytesIO()
    Image.new("1", (1, 1), 1).save(output, format="PNG")
    content = output.getvalue()
    monkeypatch.setattr("src.documents.extract_text.MAX_FILE_BYTES", len(content))
    with pytest.raises(UnsupportedDocument, match="输出|大小"):
        extract_document("resume.png", content)
