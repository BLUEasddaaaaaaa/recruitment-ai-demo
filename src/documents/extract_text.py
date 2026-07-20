"""Turn supported resume uploads into bounded text or validated image bytes."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path, PurePosixPath
from typing import Any
from zipfile import BadZipFile, LargeZipFile, ZipFile
from zlib import error as ZlibError

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from lxml.etree import XMLSyntaxError
from PIL import Image, UnidentifiedImageError
from pypdf import PdfReader
from pypdf.errors import PyPdfError
from pypdf.generic import ArrayObject, IndirectObject

MAX_FILE_BYTES = 20 * 1024 * 1024
MAX_TEXT_CHARS = 100_000
MAX_IMAGE_PIXELS = 12_000_000  # About 36 MiB when decoded as RGB; suitable for Demo concurrency.
MAX_DOCX_MEMBERS = 256
MAX_DOCX_MEMBER_BYTES = 8 * 1024 * 1024
MAX_DOCX_UNCOMPRESSED_BYTES = 40 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 100
MAX_PDF_PAGES = 50
MAX_PDF_CONTENT_STREAM_BYTES = 20 * 1024 * 1024
MAX_PDF_RESOURCE_OBJECTS = 100
MAX_PDF_RESOURCE_DEPTH = 12

_MIME_TYPES = {
    ".txt": "text/plain",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
}


class UnsupportedDocument(ValueError):
    """Raised when an upload cannot be normalized safely."""


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    filename: str
    mime_type: str
    text: str
    image_bytes: tuple[bytes, ...] = ()


class _TextAccumulator:
    def __init__(self, separator: str) -> None:
        self._separator = separator
        self._parts: list[str] = []
        self._length = 0

    def add(self, value: str) -> None:
        value = value.strip()
        if not value:
            return
        added = len(value) + (len(self._separator) if self._parts else 0)
        if self._length + added > MAX_TEXT_CHARS:
            raise UnsupportedDocument(f"提取的文本超过 {MAX_TEXT_CHARS:,} 字符限制。")
        self._parts.append(value)
        self._length += added

    def finish(self) -> str:
        return self._separator.join(self._parts)


def extract_document(filename: str, content: bytes) -> ExtractedDocument:
    """Validate and normalize one resume document without external conversion."""
    extension = Path(filename).suffix.lower()
    if extension not in _MIME_TYPES:
        raise UnsupportedDocument("不支持该文件扩展名，仅支持 PDF、DOCX、TXT、PNG 和 JPEG。")
    if not content:
        raise UnsupportedDocument("文件为空，无法处理。")
    if len(content) > MAX_FILE_BYTES:
        raise UnsupportedDocument("文件大小超过 20 MiB 限制。")
    if content.startswith((b"MZ", b"\x7fELF")) or content[:4] in {
        b"\xca\xfe\xba\xbe",
        b"\xfe\xed\xfa\xce",
        b"\xfe\xed\xfa\xcf",
        b"\xcf\xfa\xed\xfe",
    }:
        raise UnsupportedDocument("检测到可执行文件内容，已拒绝处理。")
    detected_extension = _detect_binary_extension(content)
    if detected_extension is not None and not _extensions_match(extension, detected_extension):
        raise UnsupportedDocument("文件内容与文件扩展名格式不匹配。")

    if extension == ".txt":
        result = _extract_txt(filename, content)
    elif extension == ".docx":
        _require_magic(content, b"PK\x03\x04", "DOCX")
        result = _extract_docx(filename, content)
    elif extension == ".pdf":
        _require_magic(content, b"%PDF-", "PDF")
        result = _extract_pdf(filename, content)
    else:
        result = _extract_image(filename, extension, content)

    if len(result.text) > MAX_TEXT_CHARS:
        raise UnsupportedDocument("提取的文本超过 100,000 字符限制。")
    return result


def _detect_binary_extension(content: bytes) -> str | None:
    if content.startswith(b"%PDF-"):
        return ".pdf"
    if content.startswith(b"PK\x03\x04"):
        return ".docx"
    if content.startswith(b"\x89PNG\r\n\x1a\n"):
        return ".png"
    if content.startswith(b"\xff\xd8\xff"):
        return ".jpg"
    return None


def _extensions_match(actual: str, detected: str) -> bool:
    return actual == detected or {actual, detected} <= {".jpg", ".jpeg"}


def _require_magic(content: bytes, signature: bytes, label: str) -> None:
    if not content.startswith(signature):
        raise UnsupportedDocument(f"文件内容与 {label} 扩展名格式不匹配。")


def _extract_txt(filename: str, content: bytes) -> ExtractedDocument:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError as error:
        raise UnsupportedDocument("TXT 文件必须使用有效的 UTF-8 编码。") from error
    controls = sum(character not in "\n\r\t" and ord(character) < 32 for character in text)
    if "\x00" in text or controls > max(1, len(text) // 20):
        raise UnsupportedDocument("TXT 文件包含过多控制字符，疑似二进制内容。")
    return ExtractedDocument(filename, _MIME_TYPES[".txt"], text)


def _extract_docx(filename: str, content: bytes) -> ExtractedDocument:
    try:
        _preflight_docx(content)
        document = Document(BytesIO(content))
        text = _TextAccumulator("\n")
        for paragraph in document.paragraphs:
            text.add(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.add(cell.text)
    except UnsupportedDocument:
        raise
    except (
        BadZipFile,
        LargeZipFile,
        KeyError,
        OSError,
        PackageNotFoundError,
        ValueError,
        XMLSyntaxError,
    ) as error:
        raise UnsupportedDocument("DOCX 文件已损坏或无法解析。") from error
    return ExtractedDocument(filename, _MIME_TYPES[".docx"], text.finish())


def _preflight_docx(content: bytes) -> None:
    with ZipFile(BytesIO(content)) as archive:
        members = archive.infolist()
        if len(members) > MAX_DOCX_MEMBERS:
            raise UnsupportedDocument("DOCX ZIP 条目数超过安全限制。")
        names: set[str] = set()
        total_size = 0
        for member in members:
            if member.filename in names:
                raise UnsupportedDocument("DOCX ZIP 包含重复成员名称。")
            names.add(member.filename)
            path = PurePosixPath(member.filename)
            if (
                not member.filename
                or "\\" in member.filename
                or path.is_absolute()
                or ".." in path.parts
            ):
                raise UnsupportedDocument("DOCX ZIP 包含不安全的成员路径名称。")
            if member.flag_bits & 0x1:
                raise UnsupportedDocument("DOCX ZIP 包含加密成员，无法处理。")
            if member.file_size > MAX_DOCX_MEMBER_BYTES:
                raise UnsupportedDocument("DOCX ZIP 单个成员解压后过大。")
            total_size += member.file_size
            if total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise UnsupportedDocument("DOCX ZIP 解压后总大小超过安全限制。")
            if member.file_size and member.file_size / max(member.compress_size, 1) > (
                MAX_DOCX_COMPRESSION_RATIO
            ):
                raise UnsupportedDocument("DOCX ZIP 成员压缩比异常。")
        required = {"[Content_Types].xml", "_rels/.rels", "word/document.xml"}
        if not required <= names:
            raise UnsupportedDocument("DOCX ZIP 缺少必需的 OOXML 成员。")


def _extract_pdf(filename: str, content: bytes) -> ExtractedDocument:
    try:
        reader = PdfReader(BytesIO(content), strict=True)
        if reader.is_encrypted:
            raise UnsupportedDocument("PDF 已加密或受密码保护，无法处理。")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise UnsupportedDocument(f"PDF 页数超过 {MAX_PDF_PAGES} 页限制。")
        text = _TextAccumulator("\n\n")
        stream_budget = _PdfStreamBudget()
        for page in reader.pages:
            stream_budget.visit_page(page)
            text.add(page.extract_text() or "")
    except UnsupportedDocument:
        raise
    except (PyPdfError, NotImplementedError, ValueError, TypeError, OSError, ZlibError) as error:
        raise UnsupportedDocument("PDF 文件已损坏或无法解析。") from error
    result = text.finish()
    if not result:
        raise UnsupportedDocument(
            "PDF 中未提取到文本；当前不支持将扫描页转换为图像，请先进行 OCR 或上传原始图像。"
        )
    return ExtractedDocument(filename, _MIME_TYPES[".pdf"], result)


class _PdfStreamBudget:
    """Account for page content and text-relevant Form XObject streams once."""

    def __init__(self) -> None:
        self._decoded_bytes = 0
        self._resource_objects = 0
        self._visited_streams: set[tuple[str, int, int]] = set()
        self._active_forms: set[tuple[str, int, int]] = set()

    def visit_page(self, page: Any) -> None:
        self._visit_contents(page.get("/Contents"))
        self._visit_resources(page.get_inherited("/Resources", None), depth=0)

    def _visit_contents(self, contents: Any) -> None:
        if contents is None:
            return
        resolved = contents.get_object()
        if isinstance(resolved, ArrayObject):
            for item in resolved:
                self._count_stream(item)
        else:
            self._count_stream(contents)

    def _visit_resources(self, resources: Any, depth: int) -> None:
        if resources is None:
            return
        if depth > MAX_PDF_RESOURCE_DEPTH:
            raise UnsupportedDocument("PDF 资源嵌套深度超过安全限制。")
        resources = resources.get_object()
        xobjects = resources.get("/XObject")
        if xobjects is None:
            return
        for reference in xobjects.get_object().values():
            resolved = reference.get_object()
            if resolved.get("/Subtype") != "/Form":
                # Image and other non-Form XObjects do not contain text drawing operators.
                continue
            key = self._object_key(reference, resolved)
            if key in self._active_forms:
                raise UnsupportedDocument("PDF Form 资源存在循环引用。")
            if key in self._visited_streams:
                continue
            self._resource_objects += 1
            if self._resource_objects > MAX_PDF_RESOURCE_OBJECTS:
                raise UnsupportedDocument("PDF 可达 Form 资源对象超过安全限制。")
            self._count_stream(reference)
            self._active_forms.add(key)
            try:
                self._visit_resources(resolved.get("/Resources"), depth + 1)
            finally:
                self._active_forms.remove(key)

    def _count_stream(self, reference: Any) -> None:
        resolved = reference.get_object()
        key = self._object_key(reference, resolved)
        if key in self._visited_streams:
            return
        self._visited_streams.add(key)
        self._decoded_bytes += len(resolved.get_data())
        if self._decoded_bytes > MAX_PDF_CONTENT_STREAM_BYTES:
            raise UnsupportedDocument("PDF 解码后内容流超过安全限制。")

    @staticmethod
    def _object_key(reference: Any, resolved: Any) -> tuple[str, int, int]:
        if isinstance(reference, IndirectObject):
            return ("indirect", reference.idnum, reference.generation)
        return ("direct", id(resolved), 0)


def _extract_image(filename: str, extension: str, content: bytes) -> ExtractedDocument:
    expected_format = "PNG" if extension == ".png" else "JPEG"
    signature_matches = (
        content.startswith(b"\x89PNG\r\n\x1a\n")
        if expected_format == "PNG"
        else content.startswith(b"\xff\xd8\xff")
    )
    if not signature_matches:
        raise UnsupportedDocument("图像内容与文件扩展名格式不匹配。")
    try:
        with Image.open(BytesIO(content)) as image:
            if image.format != expected_format:
                raise UnsupportedDocument("图像内容与文件扩展名格式不匹配。")
            width, height = image.size
            if width <= 0 or height <= 0 or width * height > MAX_IMAGE_PIXELS:
                raise UnsupportedDocument("图像尺寸或像素数超过安全限制。")
            image.load()
            if image.mode in {"RGBA", "LA"}:
                normalized = Image.new("RGB", image.size, "white")
                normalized.paste(image.convert("RGBA"), mask=image.getchannel("A"))
            else:
                normalized = image if image.mode == "RGB" else image.convert("RGB")
            output = BytesIO()
            normalized.save(output, format=expected_format)
    except UnsupportedDocument:
        raise
    except (UnidentifiedImageError, OSError, ValueError, Image.DecompressionBombError) as error:
        raise UnsupportedDocument("图像文件已损坏、尺寸过大或无法解码。") from error
    normalized_bytes = output.getvalue()
    if len(normalized_bytes) > MAX_FILE_BYTES:
        raise UnsupportedDocument("重新编码的图像输出大小超过安全限制。")
    return ExtractedDocument(filename, _MIME_TYPES[extension], "", (normalized_bytes,))
